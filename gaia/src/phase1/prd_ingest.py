"""Generic PRD ingestion and normalization."""
from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Sequence

from gaia.src.phase1.pdf_loader import PDFLoader
from gaia.src.phase1.prd_bundle import PRDBundle, PRDExecutionProfile, PRDFlow, PRDMetadata, PRDNormalizedDocument, PRDRequirement, PRDSection, PRDSource, is_prd_bundle_payload
from gaia.src.phase1.prd_goal_generator import generate_prd_goals

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # type: ignore[assignment]

_SECTION_HINTS = (
    "서비스 개요",
    "문제 정의",
    "타겟 사용자",
    "핵심 가치",
    "사용자 시나리오",
    "기능 요구사항",
    "정보구조",
    "비기능 요구사항",
    "데이터",
    "연동 요구사항",
    "kpi",
    "릴리즈 우선순위",
    "리스크",
)


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def _read_source_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        result = PDFLoader().extract(path)
        return result.text, "pdf"
    if suffix == ".docx":
        if Document is None:
            raise RuntimeError("python-docx가 설치되어 있지 않아 DOCX를 읽을 수 없습니다.")
        document = Document(str(path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(lines), "docx"
    if suffix in {".md", ".txt", ".prd"}:
        return path.read_text(encoding="utf-8"), suffix.lstrip(".")
    if suffix == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        if is_prd_bundle_payload(payload):
            bundle = PRDBundle.model_validate(payload)
            return bundle.model_dump_json(indent=2), "json"
    return path.read_text(encoding="utf-8"), suffix.lstrip(".") or "text"


def _detect_title(lines: Sequence[str]) -> str:
    for line in lines:
        stripped = line.strip()
        if stripped and len(stripped) <= 120:
            return stripped
    return "Untitled PRD"


def _looks_like_section_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    lowered = stripped.lower().rstrip(":")
    if any(hint in lowered for hint in _SECTION_HINTS):
        return True
    if re.match(r"^\d+[\)\.]?\s+[^\n]{1,80}$", stripped):
        return True
    return stripped.endswith(":") and len(stripped) <= 80


def _split_sections(text: str) -> List[PRDSection]:
    lines = [line.rstrip() for line in text.splitlines()]
    sections: list[PRDSection] = []
    current_title = "개요"
    current_lines: list[str] = []
    order = 1
    for line in lines:
        if _looks_like_section_heading(line):
            if current_lines:
                sections.append(PRDSection(title=current_title, content="\n".join(current_lines).strip(), order=order))
                order += 1
                current_lines = []
            current_title = line.strip().rstrip(":")
            continue
        current_lines.append(line)
    if current_lines:
        sections.append(PRDSection(title=current_title, content="\n".join(current_lines).strip(), order=order))
    return [section for section in sections if section.content.strip()]


def _extract_bullets(text: str) -> List[str]:
    items: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and re.match(r"^([•*\-]|\d+[\)\.])\s+", stripped):
            items.append(re.sub(r"^([•*\-]|\d+[\)\.])\s+", "", stripped).strip())
    return items


def _extract_requirements(sections: List[PRDSection]) -> List[PRDRequirement]:
    requirements: list[PRDRequirement] = []
    idx = 1
    for section in sections:
        title = section.title.lower()
        category = "functional"
        priority = "P1"
        if "비기능" in title or "nfr" in title:
            category = "nfr"
        if "릴리즈" in title or "우선순위" in title:
            category = "priority"
        for bullet in _extract_bullets(section.content):
            lowered = bullet.lower()
            if category == "priority":
                if "p0" in lowered or "필수" in lowered:
                    priority = "P0"
                elif "p2" in lowered or "확장" in lowered:
                    priority = "P2"
                else:
                    priority = "P1"
            requirements.append(
                PRDRequirement(
                    id=f"REQ_{idx:03d}",
                    title=bullet[:80],
                    description=bullet,
                    priority=priority,
                    category=category,
                    source_section=section.title,
                )
            )
            idx += 1
    return requirements


def _extract_flows(sections: List[PRDSection]) -> List[PRDFlow]:
    flows: list[PRDFlow] = []
    idx = 1
    for section in sections:
        title = section.title.lower()
        if "시나리오" not in title and "flow" not in title:
            continue
        steps = _extract_bullets(section.content)
        if not steps:
            continue
        flows.append(PRDFlow(id=f"FLOW_{idx:03d}", title=section.title, steps=steps, source_section=section.title))
        idx += 1
    return flows


def _extract_kpis(sections: List[PRDSection]) -> List[str]:
    rows: list[str] = []
    for section in sections:
        if "kpi" in section.title.lower():
            rows.extend(_extract_bullets(section.content))
    return rows


def _extract_risks(sections: List[PRDSection]) -> List[str]:
    rows: list[str] = []
    for section in sections:
        lowered = section.title.lower()
        if "리스크" in lowered or "위험" in lowered or "risk" in lowered:
            rows.extend(_extract_bullets(section.content))
    return rows


def _build_summary(text: str) -> str:
    sentences = [segment.strip() for segment in re.split(r"[\n\.]", text) if segment.strip()]
    return " ".join(sentences[:3])[:400]


def ingest_prd_bundle(*, input_path: str | Path | None = None, raw_text: str | None = None, base_url: str | None = None) -> PRDBundle:
    if input_path is None and not str(raw_text or "").strip():
        raise ValueError("input_path 또는 raw_text 중 하나는 필요합니다.")

    if input_path is not None:
        path = Path(input_path).expanduser()
        if not path.exists():
            raise FileNotFoundError(path)
        text, source_type = _read_source_text(path)
        source_path = str(path)
    else:
        text = str(raw_text or "")
        source_type = "text"
        source_path = None

    if source_type == "json":
        payload = json.loads(text)
        if is_prd_bundle_payload(payload):
            bundle = PRDBundle.model_validate(payload)
            if base_url:
                bundle.execution_profile.base_url = base_url
            return bundle

    lines = [line for line in text.splitlines() if line.strip()]
    title = _detect_title(lines)
    sections = _split_sections(text)
    requirements = _extract_requirements(sections)
    flows = _extract_flows(sections)
    normalized = PRDNormalizedDocument(
        summary=_build_summary(text),
        sections=sections,
        requirements=requirements,
        user_flows=flows,
        kpis=_extract_kpis(sections),
        risks=_extract_risks(sections),
    )
    now = _now_iso()
    return PRDBundle(
        project_name=title,
        source=PRDSource(type=source_type, path=source_path, content_hash=_hash_text(text), title=title),
        normalized_prd=normalized,
        generated_goals=generate_prd_goals(requirements, flows),
        execution_profile=PRDExecutionProfile(base_url=base_url),
        metadata=PRDMetadata(
            created_at=now,
            updated_at=now,
            notes=["Source normalized heuristically. Review generated goals before running."],
        ),
    )
